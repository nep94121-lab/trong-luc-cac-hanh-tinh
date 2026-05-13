#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Trích xuất nội dung từ PDF (+ chuyển ảnh JPG) và DOCX
"""

import os, sys, json
import fitz  # PyMuPDF
from docx import Document

WORKSPACE = r"C:\Users\Admin\Desktop\trọng lực ở mỗi hành tinh"
PDF_PATH  = os.path.join(WORKSPACE, "Mota trainghiem_G2.pdf")
DOCX_PATH = os.path.join(WORKSPACE, "Noi dung hoc tap.docx")
IMG_DIR   = os.path.join(WORKSPACE, "pdf_images")
OUT_FILE  = os.path.join(WORKSPACE, "noi_dung_tong_hop.md")

os.makedirs(IMG_DIR, exist_ok=True)

# ── 1. ĐỌC PDF ──────────────────────────────────────────────────────────────
print("📄 Đang xử lý PDF...")
pdf_doc = fitz.open(PDF_PATH)
pdf_text_parts = []

pdf_page_count = len(pdf_doc)
for i, page in enumerate(pdf_doc):
    # Xuất ảnh JPG độ phân giải 150 DPI
    mat = fitz.Matrix(150/72, 150/72)
    pix = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
    img_path = os.path.join(IMG_DIR, f"trang_{i+1:02d}.jpg")
    pix.save(img_path)
    print(f"  ✅ Trang {i+1}/{pdf_page_count} → {os.path.basename(img_path)}")

    # Trích text
    text = page.get_text("text").strip()
    if text:
        pdf_text_parts.append(f"### Trang {i+1}\n\n{text}")

pdf_doc.close()
pdf_full_text = "\n\n---\n\n".join(pdf_text_parts)
print(f"✅ PDF xong: {pdf_page_count} trang, ảnh lưu tại: {IMG_DIR}")

# ── 2. ĐỌC DOCX ─────────────────────────────────────────────────────────────
print("\n📝 Đang xử lý DOCX...")
docx_doc = Document(DOCX_PATH)
docx_parts = []
current_heading = None

for para in docx_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name.lower()

    if "heading 1" in style:
        docx_parts.append(f"\n## {text}")
    elif "heading 2" in style:
        docx_parts.append(f"\n### {text}")
    elif "heading 3" in style:
        docx_parts.append(f"\n#### {text}")
    elif "list" in style or para.runs and any(r.text.strip() for r in para.runs):
        # Giữ bullet nếu có
        if para.style.name.lower().startswith("list"):
            docx_parts.append(f"- {text}")
        else:
            docx_parts.append(text)
    else:
        docx_parts.append(text)

# Cũng trích bảng nếu có
table_lines = []
for tbl_idx, table in enumerate(docx_doc.tables):
    table_lines.append(f"\n**[Bảng {tbl_idx+1}]**\n")
    header = [cell.text.strip() for cell in table.rows[0].cells]
    table_lines.append("| " + " | ".join(header) + " |")
    table_lines.append("| " + " | ".join(["---"] * len(header)) + " |")
    for row in table.rows[1:]:
        cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
        table_lines.append("| " + " | ".join(cells) + " |")

docx_full_text = "\n".join(docx_parts)
if table_lines:
    docx_full_text += "\n\n---\n\n" + "\n".join(table_lines)

print(f"✅ DOCX xong: {len(docx_doc.paragraphs)} đoạn văn, {len(docx_doc.tables)} bảng")

# ── 3. TỔNG HỢP → FILE .MD ──────────────────────────────────────────────────
img_list = "\n".join(
    f"- ![Trang {i+1}](pdf_images/trang_{i+1:02d}.jpg)"
    for i in range(len(pdf_text_parts))
)

output_md = f"""# NỘI DUNG TỔNG HỢP

> File tổng hợp tự động từ:
> - `Mota trainghiem_G2.pdf`
> - `Noi dung hoc tap.docx`

---

## PHẦN 1 — MÔ TẢ TRẢI NGHIỆM G2 (PDF)

> **Ảnh các trang PDF:**
{img_list}

---

{pdf_full_text}

---

## PHẦN 2 — NỘI DUNG HỌC TẬP (DOCX)

{docx_full_text}
"""

with open(OUT_FILE, "w", encoding="utf-8") as f:
    f.write(output_md)

print(f"\n🎉 Hoàn tất! File nội dung: {OUT_FILE}")
print(f"   Ảnh JPG tại: {IMG_DIR}")
